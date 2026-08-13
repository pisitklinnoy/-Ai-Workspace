import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_report_doc():
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("รายงานการวางโครงสร้างโปรเจกต์และการสร้าง Backend APIs\n(Assignment WTN-A06)")
    title_run.font.name = "TH Sarabun PSK"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("วิชา: Web Development / API Architectureด้วย FastAPI\nหัวข้อ: Project - Lib Installations - API List")
    sub_run.font.name = "TH Sarabun PSK"
    sub_run.font.size = Pt(16)
    sub_run.font.italic = True

    doc.add_paragraph() # Spacer

    # Student Info Box
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    
    hdr_cells = info_table.rows[0].cells
    hdr_cells[0].text = "ชื่อ-นามสกุล นักศึกษา:"
    hdr_cells[1].text = "[ ระบุชื่อ-นามสกุล ของนักศึกษา ]"
    
    row2_cells = info_table.rows[1].cells
    row2_cells[0].text = "รหัสนักศึกษา:"
    row2_cells[1].text = "[ ระบุรหัสนักศึกษา ]"

    for row in info_table.rows:
        for cell in row.cells:
            set_cell_background(cell, "F2F4F7")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "TH Sarabun PSK"
                    r.font.size = Pt(14)

    doc.add_paragraph()

    def add_heading_1(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "TH Sarabun PSK"
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 51, 102)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "TH Sarabun PSK"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(51, 102, 153)
        return p

    def add_body_p(text, bold=False, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "TH Sarabun PSK"
        r.font.size = Pt(14)
        r.font.bold = bold
        r.font.italic = italic
        return p

    def add_bullet_p(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            r0 = p.add_run(bold_prefix)
            r0.font.name = "TH Sarabun PSK"
            r0.font.size = Pt(14)
            r0.font.bold = True
        r = p.add_run(text)
        r.font.name = "TH Sarabun PSK"
        r.font.size = Pt(14)
        return p

    def add_image_placeholder(label_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(6)
        set_cell_background(cell, "FAFAFA")
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"\n📸 [ แนบรูปภาพประกอบ: {label_text} ]\n(นำรูปภาพมาวางแทนที่กล่องนี้)\n")
        r.font.name = "TH Sarabun PSK"
        r.font.size = Pt(14)
        r.font.italic = True
        r.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()

    # --- SECTION 1 ---
    add_heading_1("1. การศึกษาและการวางโครงสร้างโปรเจกต์ (Project Structure)")
    add_body_p("ในการพัฒนา Backend API สำหรับระบบ Authentication ด้วย FastAPI โปรเจกต์นี้ได้รับการออกแบบตามแนวคิดสถาปัตยกรรม Separation of Concerns (SoC) หรือ Layered Architecture ซึ่งเป็นแนวทางมาตรฐานในการพัฒนาซอฟต์แวร์ระดับมืออาชีพ")
    
    add_heading_2("1.1 แนวคิดสถาปัตยกรรมที่เลือกใช้ (Layered Architecture)")
    add_body_p("แนวคิดหลักคือการแบ่งโค้ดออกเป็นชั้น ๆ (Layers) ให้แต่ละส่วนมีหน้าที่รับผิดชอบเฉพาะเจาะจง ไม่ผูกติดกันจนเกินไป (Loose Coupling) ส่งผลดีต่อระบบดังนี้:")
    add_bullet_p("ช่วยให้อ่านและทำความเข้าใจโค้ดได้ง่าย หากเกิดข้อผิดพลาดสามารถไล่หาสาเหตุใน Layer ที่เกี่ยวข้องได้ทันที", "1. ด้านการดูแลรักษา (Maintainability): ")
    add_bullet_p("สามารถขยายฟังก์ชันใหม่ ๆ เพิ่มเติมในอนาคตได้โดยไม่กระทบกระเทือนโครงสร้างเดิม", "2. ด้านการขยายระบบ (Scalability): ")
    add_bullet_p("สามารถเขียน Unit Test ทดสอบ Logic แต่ละส่วนได้อย่างเป็นอิสระ", "3. ด้านการทดสอบ (Testability): ")

    add_heading_2("1.2 อธิบายหน้าที่ของแต่ละส่วนภายในโปรเจกต์")
    
    dir_table = doc.add_table(rows=8, cols=3)
    dir_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dir_table.autofit = False

    headers = ["Directory / File", "Layer Architecture", "หน้าที่และการทำงาน"]
    for i, h in enumerate(headers):
        cell = dir_table.rows[0].cells[i]
        cell.text = h
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "TH Sarabun PSK"
                r.font.size = Pt(14)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("app/routers/", "Presentation Layer", "กำหนดเส้นทาง API (Endpoints) รับ Request และส่ง Response"),
        ("app/services/", "Business Logic Layer", "ประมวลผลตรรกะและกฎของระบบ เช่น เช็คยูสเซอร์ซ้ำ และ Hash Password"),
        ("app/schemas/", "Data Validation Layer", "ตรวจสอบรูปแบบข้อมูลขาเข้า (Input Validation) และกรองข้อมูลขาออก"),
        ("app/repositories/", "Data Access Layer", "เป็นตัวกลางในการจัดการคำสั่ง CRUD ค้นหาและบันทึกข้อมูลเข้า Database"),
        ("app/core/", "Application Layer", "โหลดการตั้งค่าระบบ (.env), ระบบความปลอดภัย Bcrypt/JWT และ Dependencies"),
        ("app/db/", "Data Layer", "เชื่อมต่อและจัดเก็บข้อมูลของแอปพลิเคชัน (Database Layer)"),
        ("app/main.py", "Entry Point", "ประกาศ FastAPI(), ตั้งค่า Metadata, รวม Router และสั่งเปิดใช้งาน Server")
    ]

    for row_idx, row_data in enumerate(data, start=1):
        row_cells = dir_table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F9FAFB")
            for p in row_cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.name = "TH Sarabun PSK"
                    r.font.size = Pt(13)

    doc.add_paragraph()
    add_image_placeholder("โครงสร้างโฟลเดอร์ของโปรเจกต์ใน IDE / Editor")

    # --- SECTION 2 ---
    add_heading_1("2. การติดตั้ง Library ที่ใช้งานในระบบ (Libraries & Dependencies)")
    add_body_p("ระบบนี้ใช้คำสั่ง pip ในการติดตั้งแพ็กเกจ Python ที่จำเป็นสำหรับการติดต่อและประมวลผล Component ต่าง ๆ โดยระบุไว้ในไฟล์ requirements.txt ดังนี้:")

    libs = [
        ("fastapi (>=0.100.0)", "Core Web Framework สำหรับสร้าง APIs ความเร็วสูง รองรับ Asynchronous และสร้าง Docs อัตโนมัติ"),
        ("uvicorn (>=0.22.0)", "ASGI Web Server ประสิทธิภาพสูง สำหรับสั่งรัน FastAPI Application"),
        ("pydantic[email] (>=2.0)", "ใช้สำหรับการตรวจสอบความถูกต้องของข้อมูล (Data Validation) และตรวจสอบอีเมล"),
        ("python-dotenv (>=1.0.0)", "ใช้สำหรับโหลด Environment Variables จากไฟล์ .env เข้าสู่ระบบ"),
        ("passlib[bcrypt] & bcrypt", "ใช้สำหรับการเข้ารหัสผ่าน (Password Hashing) ก่อนบันทึกลงฐานข้อมูลเพื่อความปลอดภัย"),
        ("pyjwt (>=2.8.0)", "ใช้สำหรับการสร้าง (Encode) และตรวจสอบ (Decode) JSON Web Token (JWT) ในการยืนยันตัวตน")
    ]

    for lib_name, lib_desc in libs:
        add_bullet_p(lib_desc, f"{lib_name}: ")

    doc.add_paragraph()

    # --- SECTION 3 ---
    add_heading_1("3. การพัฒนา Backend Server APIs (API Implementation)")
    add_body_p("จากการติดตั้ง Library ต่าง ๆ ระบบได้ถูกนำมาพัฒนาเป็น Backend Server APIs เพื่อให้บริการประมวลผลตามความต้องการของระบบ Authentication ดังนี้:")

    add_heading_2("3.1 POST /api/v1/auth/signup (ระบบสมัครสมาชิก)")
    add_bullet_p("รับข้อมูลผู้ใช้ใหม่ (username, email, full_name, password) ผ่าน Pydantic Schema (UserRegisterRequest)")
    add_bullet_p("ส่งต่อให้ AuthService ตรวจสอบว่า username หรือ email ซ้ำในระบบหรือไม่")
    add_bullet_p("ทำการ Hash รหัสผ่านด้วย Bcrypt (ผ่าน passlib) แล้วบันทึกข้อมูลลงฐานข้อมูล")
    add_bullet_p("ส่งคืน HTTP 201 Created พร้อมข้อมูลผู้ใช้ในรูปแบบ UserResponse โดยไม่ส่งรหัสผ่านออกมา")

    add_heading_2("3.2 POST /api/v1/auth/login (ระบบเข้าสู่ระบบ)")
    add_bullet_p("รับ credentials (username, password) ผ่าน UserLoginRequest")
    add_bullet_p("ตรวจสอบความถูกต้องของรหัสผ่านด้วย Bcrypt")
    add_bullet_p("เมื่อถูกต้อง ระบบจะออก JWT Access Token (กำหนดหมดอายุ) และส่งคืน HTTP 200 OK")

    add_heading_2("3.3 GET /api/v1/auth/me (ระบบดึงข้อมูลส่วนตัว - Protected Endpoint)")
    add_bullet_p("ใช้กลไก Dependency Injection (Depends(get_current_user)) ในการตรวจสอบ Header 'Authorization: Bearer <token>'")
    add_bullet_p("แกะและยืนยันความถูกต้องของ JWT Token ถ้ายืนยันสำเร็จจะส่งคืนข้อมูลโปรไฟล์ของผู้ใช้")

    doc.add_paragraph()

    # --- SECTION 4 ---
    add_heading_1("4. การจัดทำ API Documentation และการ Snapshot API List")

    add_heading_2("4.1 การใช้งาน FastAPI Metadata Arguments")
    add_body_p("ในการสร้าง API Documentation ให้มีความครบถ้วนและสะดวกต่อการนำไปใช้งาน ได้มีการศึกษาและกำหนด Arguments ต่าง ๆ ใน app/main.py และ app/routers/auth.py ดังนี้:")
    add_bullet_p("กำหนดชื่อโปรเจกต์ รายละเอียดสถาปัตยกรรม และเวอร์ชันของระบบ", "title, description, version: ")
    add_bullet_p("กำหนด URL Endpoint สำหรับเข้าชม Swagger UI และ Redoc", "docs_url, redoc_url: ")
    add_bullet_p("กำหนดคำอธิบายสั้นและคำอธิบายละเอียดของแต่ละ Endpoint", "summary, description: ")
    add_bullet_p("จัดหมวดหมู่ API ให้เป็นสัดส่วนสวยงามบนหน้าเว็บ Documentation", "tags: ")

    add_heading_2("4.2 หน้ารายงาน API Documentation (Swagger UI & Redoc)")
    add_body_p("FastAPI มีระบบสร้าง API Documentation อัตโนมัติผ่านทางเว็บเบราว์เซอร์:")
    add_bullet_p("เข้าใช้งานได้ที่ http://127.0.0.1:8000/docs สำหรับทดสอบยิง API", "1. Swagger UI: ")
    add_bullet_p("เข้าใช้งานได้ที่ http://127.0.0.1:8000/redoc สำหรับอ่านเอกสารโครงสร้าง API", "2. Redoc: ")

    add_image_placeholder("หน้าต่าง Swagger UI (http://127.0.0.1:8000/docs)")
    add_image_placeholder("หน้าต่าง Redoc (http://127.0.0.1:8000/redoc)")

    add_heading_2("4.3 วิธีการแปลง openapi.json เป็น Excel / CSV (Snapshot API List)")
    add_body_p("เพื่อทำการ Snapshot API List ของระบบไว้ใช้ในการอ้างอิง จึงได้เลือกใช้วิธีการเขียน Custom Python Script ในไฟล์ export_api_list.py มีขั้นตอนการทำงานดังนี้:")
    add_bullet_p("ดึงโครงสร้าง OpenAPI Schema จาก app.openapi() ของ FastAPI และบันทึกเป็นไฟล์ openapi.json", "ขั้นตอนที่ 1: ")
    add_bullet_p("อ่านข้อมูล JSON และวนลูปสกัดฟิลด์สำคัญ ได้แก่ HTTP Method, Path, Summary, Tags, Request Body, Response Codes และ Description", "ขั้นตอนที่ 2: ")
    add_bullet_p("เขียนข้อมูลทั้งหมดลงไฟล์ api_list.csv ( encode ด้วย UTF-8 with BOM) ทำให้เปิดด้วย Microsoft Excel ได้โดยภาษาไทยไม่ต่าง", "ขั้นตอนที่ 3: ")

    add_heading_2("4.4 ตารางผลลัพธ์ Snapshot API List (ตารางจาก api_list.csv)")

    api_table = doc.add_table(rows=5, cols=6)
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_table.autofit = False

    t_headers = ["Method", "Path", "Summary", "Tags", "Request Body", "Response Codes"]
    for i, h in enumerate(t_headers):
        cell = api_table.rows[0].cells[i]
        cell.text = h
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "TH Sarabun PSK"
                r.font.size = Pt(13)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    api_rows = [
        ("POST", "/api/v1/auth/signup", "User Signup / Registration", "Authentication API", "Yes", "201, 422"),
        ("POST", "/api/v1/auth/login", "User Login / Signin", "Authentication API", "Yes", "200, 422"),
        ("GET", "/api/v1/auth/me", "Get Authenticated User Profile", "Authentication API", "No", "200"),
        ("GET", "/", "Root Health Check", "Health Check", "No", "200")
    ]

    for row_idx, row_data in enumerate(api_rows, start=1):
        row_cells = api_table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F9FAFB")
            for p in row_cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.name = "TH Sarabun PSK"
                    r.font.size = Pt(12)

    doc.add_paragraph()
    add_image_placeholder("ผลลัพธ์การเปิดไฟล์ api_list.csv บน Microsoft Excel")

    # --- SECTION 5 ---
    add_heading_1("5. การแนบ README.md ใน Directory สำคัญ")
    add_body_p("เพื่ออธิบายการทำงานและโน้ตสำคัญของแต่ละ Component ในโปรเจกต์ ได้มีการสร้างไฟล์ README.md กำกับไว้ในทุกโฟลเดอร์หลัก ดังนี้:")

    add_bullet_p("อธิบายหน้าที่การรับส่ง HTTP Request/Response ของ Endpoints", "app/routers/README.md: ")
    add_bullet_p("อธิบายกฎตรรกะทางธุรกิจและการประมวลผลหลักของระบบ", "app/services/README.md: ")
    add_bullet_p("อธิบายการกำหนด Schema สำหรับ Input Validation และ Output Masking", "app/schemas/README.md: ")
    add_bullet_p("อธิบายการจัดการคำสั่ง CRUD และการเชื่อมโยงกับ Database", "app/repositories/README.md: ")
    add_bullet_p("อธิบายระบบ Config, Security (Bcrypt/JWT) และ Dependencies", "app/core/README.md: ")
    add_bullet_p("อธิบายการจัดเก็บข้อมูลและเอนจินฐานข้อมูล", "app/db/README.md: ")

    doc.add_paragraph()

    # Save document
    filename = "WTN-A06_Report_FastAPI.docx"
    doc.save(filename)
    print(f"Docx generated successfully: {filename}")

if __name__ == "__main__":
    create_report_doc()
